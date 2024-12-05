import re
from typing import List, Dict, Optional, Tuple
import os
from pathlib import Path
from typing import List, Dict, Optional
import hashlib
import json
from datetime import datetime

# Document processing
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize

# Vector store
from chromadb import PersistentClient, Settings
import chromadb

# Embedding model
from sentence_transformers import SentenceTransformer

# dotenv
from dotenv import load_dotenv
load_dotenv()

# LLM
from groq import Groq

def initialize_nltk():
    """Download required NLTK data"""
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')

initialize_nltk()

class DocumentProcessor:
    def __init__(self, docs_dir: str, chunk_size: int = 50000, chunk_overlap: int = 30000):
        self.docs_dir = Path(docs_dir)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def get_document_hash(self, file_path: Path) -> str:
        """Generate MD5 hash of document for change detection"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    
    def extract_page_numbers(self, doc) -> Dict[int, int]:
        """Extract page numbers for paragraphs"""
        page_map = {}
        current_page = 1
        
        for idx, paragraph in enumerate(doc.paragraphs):
            # Check for page breaks
            for run in paragraph.runs:
                if '<w:br w:type="page"/>' in run._element.xml:
                    current_page += 1
                    
            page_map[idx] = current_page
            
        return page_map

    def find_article_boundaries(self, doc) -> List[Dict]:
        """Find article boundaries in the document with page numbers"""
        articles = []
        current_article = []
        current_article_number = None
        start_page = 1
        end_page = 1
        page_map = self.extract_page_numbers(doc)
        
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            current_page = page_map.get(idx, 1)
            
            # Check for new article
            article_match = re.match(r'(?i)^\s*article\s+(\d+)[:\.]?\s*', text)
            if article_match:
                # Save previous article if exists
                if current_article:
                    articles.append({
                        'number': current_article_number,
                        'content': current_article,
                        'pages': {
                            'start': start_page,
                            'end': end_page
                        }
                    })
                current_article = []
                current_article_number = article_match.group(1)
                start_page = current_page
            
            current_article.append(text)
            end_page = current_page
        
        # Add the last article
        if current_article:
            articles.append({
                'number': current_article_number,
                'content': current_article,
                'pages': {
                    'start': start_page,
                    'end': end_page
                }
            })
        
        return articles


    def detect_section_type(self, text: str, style_name: str) -> Dict:
        """Detect the type and details of a section"""
        text = text.strip()
        if not text:
            return None
            
        # Check for various section types
        article_match = re.match(r'(?i)^\s*article\s+(\d+)[:\.]?\s*', text)
        toc_match = re.match(r'(?i)^\s*(table\s+of\s+contents|contents)', text)
        appendix_match = re.match(r'(?i)^\s*(appendix|annex)\s+([A-Za-z0-9])[:\.]?\s*', text)
        section_match = re.match(r'(?i)^\s*(section|chapter)\s+(\d+)[:\.]?\s*', text)
        
        if article_match:
            return {
                'type': 'article',
                'number': article_match.group(1),
                'title': text
            }
        elif toc_match:
            return {
                'type': 'toc',
                'title': text
            }
        elif appendix_match:
            return {
                'type': 'appendix',
                'number': appendix_match.group(2),
                'title': text
            }
        elif section_match:
            return {
                'type': 'section',
                'number': section_match.group(2),
                'title': text
            }
        elif style_name.startswith('Heading'):
            return {
                'type': 'heading',
                'title': text
            }
            
        return None

    def extract_document_structure(self, doc) -> List[Dict]:
        """Extract all document content with structure preservation"""
        sections = []
        current_section = []
        current_metadata = None
        document_stats = {
            'articles': set(),
            'sections': set(),
            'appendices': set()
        }
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Detect section type
            section_info = self.detect_section_type(text, paragraph.style.name)
            
            if section_info:
                # Save previous section if exists
                if current_section:
                    sections.append({
                        'metadata': current_metadata,
                        'content': current_section
                    })
                
                current_section = []
                current_metadata = section_info
                
                # Update document statistics
                if section_info['type'] == 'article':
                    document_stats['articles'].add(section_info['number'])
                elif section_info['type'] == 'section':
                    document_stats['sections'].add(section_info['number'])
                elif section_info['type'] == 'appendix':
                    document_stats['appendices'].add(section_info['number'])
            
            current_section.append(text)
        
        # Add the last section
        if current_section:
            sections.append({
                'metadata': current_metadata,
                'content': current_section
            })
        
        # Add document statistics as a special section
        stats_text = [
            "Document Statistics:",
            f"Total Articles: {len(document_stats['articles'])}",
            f"Total Sections: {len(document_stats['sections'])}",
            f"Total Appendices: {len(document_stats['appendices'])}"
        ]
        
        sections.append({
            'metadata': {'type': 'statistics', 'title': 'Document Statistics'},
            'content': stats_text
        })
        
        return sections

    def process_docx(self, file_path: Path) -> List[Dict]:
        """Process document with comprehensive content handling"""
        doc = Document(file_path)
        sections = self.extract_document_structure(doc)
        
        chunks = []
        for section in sections:
            if not section['metadata']:
                continue
                
            # Keep section content together
            section_text = "\n".join(section['content'])
            
            # Create metadata
            metadata = {
                'source': str(file_path),
                'type': section['metadata']['type'],
                'title': section['metadata']['title'],
                'timestamp': datetime.now().isoformat()
            }
            
            # Add specific metadata based on type
            if section['metadata']['type'] in ['article', 'section', 'appendix']:
                metadata['number'] = section['metadata']['number']
            
            chunks.append({
                'text': section_text,
                'metadata': metadata
            })
        
        return chunks

class RAGSystem:
    def __init__(self, docs_dir: str, db_dir: str):
        self.docs_dir = Path(docs_dir)
        self.db_dir = Path(db_dir)
        self.processor = DocumentProcessor(docs_dir)
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Ensure clean directory
        if os.path.exists(str(self.db_dir)):
            import shutil
            shutil.rmtree(str(self.db_dir))
        os.makedirs(str(self.db_dir), exist_ok=True)
        
        # Initialize ChromaDB with specific settings
        settings = Settings(
            is_persistent=True,
            persist_directory=str(self.db_dir),
            anonymized_telemetry=False
        )
        
        try:
            self.chroma_client = PersistentClient(
                path=str(self.db_dir),
                settings=settings
            )
            
            # Create new collection
            self.collection = self.chroma_client.create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
        except Exception as e:
            print(f"Error initializing ChromaDB: {e}")
            raise
        
        # Initialize Groq client
        self.groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        
        # Load or initialize document hashes
        self.hashes_file = self.db_dir / "document_hashes.json"
        self.document_hashes = self._load_document_hashes()

    def _load_document_hashes(self) -> Dict[str, str]:
        if self.hashes_file.exists():
            with open(self.hashes_file, 'r') as f:
                return json.load(f)
        return {}
    
    def _save_document_hashes(self):
        with open(self.hashes_file, 'w') as f:
            json.dump(self.document_hashes, f)


    def update_knowledge_base(self):
        current_files = list(self.docs_dir.glob("*.docx"))
        
        for file_path in current_files:
            current_hash = self.processor.get_document_hash(file_path)
            if str(file_path) not in self.document_hashes or self.document_hashes[str(file_path)] != current_hash:
                print(f"Processing updated file: {file_path}")
                
                if str(file_path) in self.document_hashes:
                    self.collection.delete(where={"source": str(file_path)})
                
                chunks = self.processor.process_docx(file_path)
                
                if not chunks:
                    print(f"Warning: No content found in {file_path}")
                    continue
                
                texts = [chunk['text'] for chunk in chunks]
                embeddings = self.embedding_model.encode(texts).tolist()
                metadatas = [chunk['metadata'] for chunk in chunks]
                
                # Generate unique IDs using content hash
                ids = []
                used_ids = set()  # Track used IDs to ensure uniqueness
                
                for idx, chunk in enumerate(chunks):
                    metadata = chunk['metadata']
                    content_type = metadata.get('type', 'general')
                    content_hash = hashlib.md5(chunk['text'].encode()).hexdigest()[:8]
                    
                    # Create base ID
                    if content_type in ['article', 'section', 'appendix'] and 'number' in metadata:
                        base_id = f"{file_path}_{content_type}_{metadata['number']}"
                    else:
                        base_id = f"{file_path}_{content_type}_{idx}"
                    
                    # Ensure ID uniqueness by adding content hash
                    id_str = f"{base_id}_{content_hash}"
                    
                    # In the unlikely event of a collision, append an index
                    counter = 1
                    while id_str in used_ids:
                        id_str = f"{base_id}_{content_hash}_{counter}"
                        counter += 1
                    
                    used_ids.add(id_str)
                    ids.append(id_str)
                
                # Add chunks in smaller batches to handle large documents
                batch_size = 100
                for i in range(0, len(texts), batch_size):
                    end_idx = min(i + batch_size, len(texts))
                    self.collection.add(
                        embeddings=embeddings[i:end_idx],
                        documents=texts[i:end_idx],
                        metadatas=metadatas[i:end_idx],
                        ids=ids[i:end_idx]
                    )
                
                self.document_hashes[str(file_path)] = current_hash
        
        self._save_document_hashes()

    def query(self, question: str, n_results: int = 3) -> Dict:
        # Analyze question for specific content requests
        article_match = re.search(r'(?i)article\s+(\d+)', question)
        toc_match = re.search(r'(?i)table\s+of\s+contents|toc', question)
        stats_match = re.search(r'(?i)how\s+many|total|count|number\s+of', question)
        
        # Generate embedding for the question
        question_embedding = self.embedding_model.encode(question).tolist()
        
        try:
            # Handle different types of queries
            if article_match:
                results = self.collection.query(
                    query_embeddings=[question_embedding],
                    where={"$and": [
                        {"type": "article"},
                        {"number": article_match.group(1)}
                    ]},
                    n_results=1
                )
            elif toc_match:
                results = self.collection.query(
                    query_embeddings=[question_embedding],
                    where={"type": "toc"},
                    n_results=1
                )
            elif stats_match:
                results = self.collection.query(
                    query_embeddings=[question_embedding],
                    where={"type": "statistics"},
                    n_results=1
                )
            else:
                results = self.collection.query(
                    query_embeddings=[question_embedding],
                    n_results=n_results
                )

            if not results['ids'][0]:
                return {
                    'answer': "I couldn't find the requested information in the document.",
                    'sources': []
                }

            # Build context from results
            context_parts = []
            for doc, meta in zip(results['documents'][0], results['metadatas'][0]):
                section_type = meta.get('type', 'Section')
                title = meta.get('title', '')
                if section_type in ['article', 'section', 'appendix']:
                    section_info = f"{section_type.title()} {meta.get('number', '')}: {title}"
                else:
                    section_info = title
                
                context_parts.append(f"Source: {meta['source']}\n{section_info}:\n{doc}")

            context = "\n\n---\n\n".join(context_parts)
            
            prompt = f"""You are a helpful assistant explaining content from a document. Please follow these instructions carefully:

1. Provide accurate answers based on the context provided.
2. For articles, sections, or appendices, include their complete text first, then explain if needed.
3. For statistics or counts, ensure accuracy and cite the specific numbers found in the document.
4. Always include relevant section numbers and titles in your response.
5. If the requested information isn't in the context, say so clearly.

Context:
{context}

Question: {question}

Answer:"""

            response = self.groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="mixtral-8x7b-32768",
                temperature=0.1,
            )

            # Format sources
            sources = []
            for meta in results['metadatas'][0]:
                section_type = meta.get('type', 'Section').title()
                if section_type in ['Article', 'Section', 'Appendix']:
                    source = f"{meta['source']} ({section_type} {meta.get('number', '')})"
                else:
                    source = f"{meta['source']} ({meta.get('title', 'General Content')})"
                sources.append(source)

            return {
                'answer': response.choices[0].message.content,
                'sources': sources
            }
            
        except Exception as e:
            print(f"Error during query: {e}")
            return {
                'answer': "I encountered an error while searching. Please try again.",
                'sources': []
            }

    def get_feedback(self, question: str, answer: str, feedback: str):
        feedback_file = self.db_dir / "feedback.jsonl"
        feedback_entry = {
            'timestamp': datetime.now().isoformat(),
            'question': question,
            'answer': answer,
            'feedback': feedback
        }
        
        with open(feedback_file, 'a') as f:
            f.write(json.dumps(feedback_entry) + '\n')


if __name__ == "__main__":
    # Create necessary directories if they don't exist
    DOCS_DIR = os.getenv("DOCS_DIR")
    if not DOCS_DIR:
        raise ValueError("DOCS_DIR environment variable not set")
    os.makedirs(DOCS_DIR, exist_ok=True)
    
    # Clean start
    if os.path.exists("./rag_db"):
        import shutil
        print("Cleaning up existing database...")
        shutil.rmtree("./rag_db")
    
    os.makedirs("./rag_db", exist_ok=True)
    
    print("Initializing RAG system...")
    try:
        rag = RAGSystem(
            docs_dir=DOCS_DIR,
            db_dir="./rag_db"
        )
        
        print("Updating knowledge base...")
        rag.update_knowledge_base()
        
        print("Ready for queries!")
        
        # Interactive query loop
        while True:
            question = input("\nEnter your question (or 'quit' to exit): ")
            if question.lower() == 'quit':
                break
                
            result = rag.query(question)
            print("\nAnswer:", result['answer'])
            print("\nSources:", result['sources'])
            
            feedback = input("\nWas this answer helpful? (yes/no): ")
            rag.get_feedback(question, result['answer'], feedback)
            
    except Exception as e:
        print(f"Error: {e}")
        print("Cleaning up...")
        if os.path.exists("./rag_db"):
            import shutil
            shutil.rmtree("./rag_db")
        raise